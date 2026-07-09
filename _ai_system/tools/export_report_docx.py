from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import shutil
import subprocess
import sys
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup, Comment, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path("00_사용자_작업공간")
DEFAULT_OUTPUT = Path("reports") / "internal_review_report.docx"
LIST_STYLE_PRESETS_PATH = Path("_ai_system") / "document_presets" / "list_style_presets.json"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 0

FONT_BODY = "Malgun Gothic"
FONT_MONO = "Consolas"

COLOR_DARK = "172033"
COLOR_INK = "1F2933"
COLOR_MUTED = "616670"
COLOR_LINE = "CFD0D3"
COLOR_SOFT = "EDF3FA"
COLOR_BLUE = "2563EB"
COLOR_WARNING = "FFF7ED"

FALLBACK_LIST_STYLE_PRESETS: dict[str, list[dict[str, str]]] = {
    "formal_outline": [
        {"docx_numFmt": "upperRoman", "docx_level_text": "%1"},
        {"docx_numFmt": "upperLetter", "docx_level_text": "%2"},
        {"docx_numFmt": "decimal", "docx_level_text": "%3"},
        {"docx_numFmt": "lowerLetter", "docx_level_text": "%4"},
    ],
    "guide_outline": [
        {"docx_numFmt": "upperLetter", "docx_level_text": "%1"},
        {"docx_numFmt": "upperLetter", "docx_level_text": "%2)"},
        {"docx_numFmt": "lowerLetter", "docx_level_text": "%3)"},
        {"docx_numFmt": "lowerLetter", "docx_level_text": "(%4)"},
    ],
    "procedure_steps": [
        {"docx_numFmt": "decimal", "docx_level_text": "%1"},
        {"docx_numFmt": "decimal", "docx_level_text": "%2)"},
        {"docx_numFmt": "lowerLetter", "docx_level_text": "%3)"},
        {"docx_numFmt": "lowerLetter", "docx_level_text": "(%4)"},
    ],
    "administrative_outline": [
        {"docx_numFmt": "decimal", "docx_level_text": "%1."},
        {"docx_numFmt": "decimal", "docx_level_text": "%2)"},
        {"docx_numFmt": "upperLetter", "docx_level_text": "%3."},
        {"docx_numFmt": "lowerLetter", "docx_level_text": "%4)"},
    ],
    "symbol_bullets": [
        {"docx_numFmt": "bullet", "docx_level_text": "•"},
        {"docx_numFmt": "bullet", "docx_level_text": "◦"},
        {"docx_numFmt": "bullet", "docx_level_text": "▪"},
        {"docx_numFmt": "bullet", "docx_level_text": "-"},
    ],
}


def now_kst() -> str:
    return datetime.now(timezone(timedelta(hours=9))).isoformat(timespec="seconds")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="strict")


def write_json(path: Path, payload: dict[str, object]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8", newline="\n")


def sha256_file(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def hex_color(value: str) -> RGBColor:
    value = value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(
    run,
    *,
    name: str = FONT_BODY,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = hex_color(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_paragraph_format(paragraph, *, before: float = 0, after: float = 6, line: float = 1.18) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.strip("#"))


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = COLOR_LINE, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color.strip("#"))


def set_table_width(table, width_dxa: int = CONTENT_WIDTH_DXA, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))
    cell.width = Inches(width_dxa / 1440)


def paragraph_border_bottom(paragraph, color: str = COLOR_LINE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color.strip("#"))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = hex_color(COLOR_INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 18, COLOR_DARK, 18, 8),
        ("Heading 2", 14, COLOR_DARK, 14, 6),
        ("Heading 3", 12.5, COLOR_DARK, 10, 4),
    ]:
        style = styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = hex_color(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def load_list_style_presets() -> dict[str, list[dict[str, str]]]:
    if not LIST_STYLE_PRESETS_PATH.exists():
        return FALLBACK_LIST_STYLE_PRESETS
    try:
        payload = json.loads(read_text(LIST_STYLE_PRESETS_PATH))
    except (json.JSONDecodeError, OSError):
        return FALLBACK_LIST_STYLE_PRESETS
    presets: dict[str, list[dict[str, str]]] = {}
    for preset in payload.get("presets", []):
        if not isinstance(preset, dict):
            continue
        preset_id = str(preset.get("preset_id", "")).strip()
        levels = preset.get("levels", [])
        if not preset_id or not isinstance(levels, list):
            continue
        cleaned: list[dict[str, str]] = []
        for level in levels[:4]:
            if not isinstance(level, dict):
                continue
            cleaned.append(
                {
                    "docx_numFmt": str(level.get("docx_numFmt", "")).strip() or "decimal",
                    "docx_level_text": str(level.get("docx_level_text", "")).strip() or "%1",
                }
            )
        if len(cleaned) == 4:
            presets[preset_id] = cleaned
    return presets or FALLBACK_LIST_STYLE_PRESETS


def _next_numbering_id(numbering, tag: str, attr: str) -> int:
    values: list[int] = []
    for child in numbering.findall(qn(f"w:{tag}")):
        value = child.get(qn(f"w:{attr}"))
        if value is not None and str(value).isdigit():
            values.append(int(value))
    return max(values, default=-1) + 1


def _append_val(parent, tag: str, value: str) -> None:
    node = OxmlElement(tag)
    node.set(qn("w:val"), str(value))
    parent.append(node)


def _append_indentation(parent, *, left: int, hanging: int) -> None:
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    parent.append(p_pr)


def configure_list_numbering(doc: Document, presets: dict[str, list[dict[str, str]]]) -> dict[str, int]:
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "abstractNum", "abstractNumId")
    num_id = _next_numbering_id(numbering, "num", "numId")
    preset_num_ids: dict[str, int] = {}

    for preset_id, levels in presets.items():
        current_abstract_id = abstract_id
        current_num_id = num_id
        abstract_id += 1
        num_id += 1

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(current_abstract_id))
        _append_val(abstract, "w:multiLevelType", "hybridMultilevel")

        for index, level in enumerate(levels[:4]):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(index))
            _append_val(lvl, "w:start", "1")
            _append_val(lvl, "w:numFmt", level["docx_numFmt"])
            _append_val(lvl, "w:lvlText", level["docx_level_text"])
            _append_val(lvl, "w:lvlJc", "left")
            _append_indentation(lvl, left=720 + index * 360, hanging=260)
            if level["docx_numFmt"] == "bullet":
                r_pr = OxmlElement("w:rPr")
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), FONT_BODY)
                r_fonts.set(qn("w:hAnsi"), FONT_BODY)
                r_fonts.set(qn("w:eastAsia"), FONT_BODY)
                r_pr.append(r_fonts)
                lvl.append(r_pr)
            abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(current_num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(current_abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)
        preset_num_ids[preset_id] = current_num_id

    return preset_num_ids


def set_paragraph_numbering(paragraph, *, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), str(max(0, min(level, 3))))
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def set_running_header_footer(doc: Document, cover_data: dict[str, object]) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(cover_data.get("report_title", "")).strip() or "Report")
    set_run_font(run, size=8.5, color=COLOR_MUTED, bold=True)
    if cover_data.get("version"):
        run = p.add_run(f"  |  {cover_data['version']}")
        set_run_font(run, size=8.5, color=COLOR_MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = str(cover_data.get("report_no", "")).strip() or "Native DOCX export"
    run = p.add_run(label)
    set_run_font(run, size=8, color=COLOR_MUTED)


def add_text(paragraph, text: str, *, size: float = 10.5, color: str = COLOR_INK, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_cover(doc: Document, cover_data: dict[str, object]) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, after=6, line=1.0)
    if cover_data.get("classification"):
        run = p.add_run(str(cover_data["classification"]).strip())
        set_run_font(run, size=9.5, color=COLOR_DARK, bold=True)
    if cover_data.get("security_tag"):
        run = p.add_run("   " + str(cover_data["security_tag"]).strip())
        set_run_font(run, size=9.5, color=COLOR_MUTED, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=10, after=3, line=1.0)
    run = p.add_run(str(cover_data.get("kicker", "")).strip())
    set_run_font(run, size=10.5, color=COLOR_BLUE, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=2, after=8, line=1.12)
    run = p.add_run(str(cover_data.get("report_title", "")).strip() or "보고서")
    set_run_font(run, size=24, color=COLOR_DARK, bold=True)

    subtitle = str(cover_data.get("subtitle", "")).strip()
    if subtitle:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=14, line=1.25)
        add_text(p, subtitle, size=11.5, color=COLOR_MUTED)

    rows = [
        ("프로젝트", cover_data.get("project_name", ""), "보고서 번호", cover_data.get("report_no", "")),
        ("작성일", cover_data.get("date", ""), "버전", cover_data.get("version", "")),
        ("작성 주체", cover_data.get("prepared_by", ""), "보안 등급", cover_data.get("security_level", "")),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_width(table)
    set_table_borders(table)
    widths = [1500, 3180, 1500, 3180]
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            set_paragraph_format(p, after=0, line=1.1)
            if col_index in {0, 2}:
                set_cell_shading(cell, COLOR_SOFT)
                add_text(p, str(value), size=9.2, color=COLOR_DARK, bold=True)
            else:
                add_text(p, str(value), size=9.2, color=COLOR_INK)

    doc.add_paragraph()
    approvals = [
        ("작성", cover_data.get("approval_author", "")),
        ("검토", cover_data.get("approval_reviewer", "")),
        ("승인", cover_data.get("approval_approver", "")),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table)
    set_table_borders(table, color=COLOR_LINE)
    for index, (label, name) in enumerate(approvals):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, CONTENT_WIDTH_DXA // 3)
        set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
        p = cell.paragraphs[0]
        set_paragraph_format(p, after=4, line=1.1)
        add_text(p, label, size=8.8, color=COLOR_MUTED)
        p = cell.add_paragraph()
        set_paragraph_format(p, after=0, line=1.1)
        add_text(p, str(name), size=10.2, color=COLOR_DARK, bold=True)

    purpose = str(cover_data.get("purpose", "")).strip()
    if purpose:
        doc.add_paragraph()
        add_callout(doc, purpose, fill=COLOR_SOFT)

    doc.add_page_break()


def iter_meaningful_children(node: Tag) -> Iterable[Tag | NavigableString]:
    for child in node.children:
        if isinstance(child, Comment):
            continue
        if isinstance(child, NavigableString):
            if str(child).strip():
                yield child
            continue
        if isinstance(child, Tag):
            yield child


def text_content(node: Tag | NavigableString) -> str:
    if isinstance(node, NavigableString):
        return str(node)
    return " ".join(node.get_text(" ", strip=True).split())


def add_inline(paragraph, node: Tag | NavigableString, *, base_size: float = 10.5) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=base_size, color=COLOR_INK)
        return

    tag = node.name.lower()
    if tag == "br":
        paragraph.add_run().add_break()
        return
    if tag in {"strong", "b"}:
        for child in node.children:
            if isinstance(child, NavigableString):
                run = paragraph.add_run(str(child))
                set_run_font(run, size=base_size, color=COLOR_INK, bold=True)
            else:
                add_inline(paragraph, child, base_size=base_size)
        return
    if tag in {"em", "i"}:
        for child in node.children:
            if isinstance(child, NavigableString):
                run = paragraph.add_run(str(child))
                set_run_font(run, size=base_size, color=COLOR_INK, italic=True)
            else:
                add_inline(paragraph, child, base_size=base_size)
        return
    if tag == "code":
        run = paragraph.add_run(text_content(node))
        set_run_font(run, name=FONT_MONO, size=max(base_size - 0.5, 8), color=COLOR_DARK)
        return
    if tag == "a":
        run = paragraph.add_run(text_content(node))
        set_run_font(run, size=base_size, color=COLOR_BLUE, underline=True)
        return
    for child in node.children:
        add_inline(paragraph, child, base_size=base_size)


def add_paragraph_from_node(doc: Document, node: Tag, *, style: str | None = None, size: float = 10.5) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, after=6, line=1.18)
    add_text(p, text_content(node), size=size, color=COLOR_INK)
    if not p.text.strip():
        p._element.getparent().remove(p._element)


def add_heading(doc: Document, node: Tag, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    if level == 2:
        paragraph_border_bottom(p)
    run = p.add_run(text_content(node))
    size = {1: 18, 2: 14, 3: 12.5}.get(level, 12)
    set_run_font(run, size=size, color=COLOR_DARK, bold=True)


def add_callout(doc: Document, text: str, fill: str = "FFFFFF") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, color=COLOR_BLUE, size="8")
    cell = table.rows[0].cells[0]
    set_cell_width(cell, CONTENT_WIDTH_DXA)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=0, line=1.18)
    add_text(p, text, size=10, color=COLOR_INK)
    doc.add_paragraph()


def explicit_list_preset(node: Tag, known_preset_ids: set[str]) -> str:
    direct = str(node.get("data-list-preset", "") or node.get("data-list-style", "")).strip()
    if direct in known_preset_ids:
        return direct
    class_tokens = {str(item).replace("-", "_") for item in node.get("class", [])}
    for preset_id in known_preset_ids:
        if preset_id in class_tokens or f"list_preset_{preset_id}" in class_tokens:
            return preset_id
    return ""


def choose_list_preset(
    node: Tag,
    *,
    ordered: bool,
    known_preset_ids: set[str],
    inherited_preset_id: str | None = None,
) -> str:
    explicit = explicit_list_preset(node, known_preset_ids)
    if explicit:
        return explicit
    if ordered and inherited_preset_id in known_preset_ids:
        return str(inherited_preset_id)
    if not ordered and "symbol_bullets" in known_preset_ids:
        return "symbol_bullets"
    if ordered and "formal_outline" in known_preset_ids:
        return "formal_outline"
    return next(iter(known_preset_ids), "")


def iter_list_item_inline_children(li: Tag) -> Iterable[Tag | NavigableString]:
    for child in li.children:
        if isinstance(child, Comment):
            continue
        if isinstance(child, NavigableString):
            if str(child).strip():
                yield child
            continue
        if isinstance(child, Tag):
            if child.name.lower() in {"ol", "ul"}:
                continue
            yield child


def iter_direct_child_lists(li: Tag) -> Iterable[Tag]:
    for child in li.children:
        if isinstance(child, Tag) and child.name.lower() in {"ol", "ul"}:
            yield child


def add_list(
    doc: Document,
    node: Tag,
    ordered: bool,
    numbering_ids: dict[str, int],
    *,
    level: int = 0,
    inherited_preset_id: str | None = None,
) -> None:
    known_preset_ids = set(numbering_ids)
    preset_id = choose_list_preset(
        node,
        ordered=ordered,
        known_preset_ids=known_preset_ids,
        inherited_preset_id=inherited_preset_id,
    )
    fallback_style = "List Number" if ordered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        p = doc.add_paragraph(style=None if preset_id else fallback_style)
        set_paragraph_format(p, after=4, line=1.15)
        if preset_id:
            set_paragraph_numbering(p, num_id=numbering_ids[preset_id], level=level)
        else:
            p.paragraph_format.left_indent = Inches(0.42 + min(level, 3) * 0.22)
            p.paragraph_format.first_line_indent = Inches(-0.18)
        for child in iter_list_item_inline_children(li):
            add_inline(p, child, base_size=10.3)
        if not p.text.strip():
            p._element.getparent().remove(p._element)
        for nested in iter_direct_child_lists(li):
            nested_ordered = nested.name.lower() == "ol"
            nested_inherited = preset_id if nested_ordered else "symbol_bullets"
            add_list(
                doc,
                nested,
                ordered=nested_ordered,
                numbering_ids=numbering_ids,
                level=min(level + 1, 3),
                inherited_preset_id=nested_inherited,
            )


def add_caption(doc: Document, text: str, *, figure: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=3, after=5 if figure else 3, line=1.12)
    add_text(p, text, size=9.2, color=COLOR_MUTED, bold=figure)


def add_table(doc: Document, node: Tag) -> None:
    caption = node.find("caption", recursive=False)
    if caption:
        add_caption(doc, text_content(caption), figure=False)

    rows: list[list[tuple[str, bool]]] = []
    for tr in node.find_all("tr"):
        cells: list[tuple[str, bool]] = []
        for cell in tr.find_all(["th", "td"], recursive=False):
            cells.append((text_content(cell), cell.name.lower() == "th"))
        if cells:
            rows.append(cells)
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_width(table)
    set_table_borders(table)
    width = CONTENT_WIDTH_DXA // cols
    for r_index, row in enumerate(rows):
        for c_index in range(cols):
            cell = table.rows[r_index].cells[c_index]
            set_cell_width(cell, width)
            set_cell_margins(cell, top=115, bottom=115, start=140, end=140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text, is_header = row[c_index] if c_index < len(row) else ("", False)
            if is_header or r_index == 0 and any(item[1] for item in row):
                set_cell_shading(cell, COLOR_SOFT)
            p = cell.paragraphs[0]
            set_paragraph_format(p, after=0, line=1.12)
            add_text(p, text, size=9.0, color=COLOR_DARK if is_header else COLOR_INK, bold=is_header)
    doc.add_paragraph()


def resolve_image(project: Path, src: str) -> Path | None:
    src = src.strip()
    if not src:
        return None
    candidates = []
    raw = Path(src)
    if raw.is_absolute():
        candidates.append(raw)
    candidates.extend([
        project / "reports" / src,
        project / src,
    ])
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def add_figure(doc: Document, project: Path, node: Tag) -> None:
    img = node.find("img")
    if img and img.get("src"):
        image_path = resolve_image(project, str(img.get("src")))
        if image_path:
            p = doc.add_paragraph()
            set_paragraph_format(p, after=4, line=1.0)
            p.paragraph_format.keep_with_next = True
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(5.4))
    caption = node.find("figcaption")
    if caption:
        add_caption(doc, text_content(caption), figure=True)
    doc.add_paragraph()


def render_node(doc: Document, project: Path, node: Tag | NavigableString, numbering_ids: dict[str, int]) -> None:
    if isinstance(node, NavigableString):
        if str(node).strip():
            p = doc.add_paragraph()
            add_text(p, str(node).strip())
        return
    tag = node.name.lower()
    if tag in {"script", "style"}:
        return
    if tag == "h1":
        add_heading(doc, node, 1)
    elif tag == "h2":
        add_heading(doc, node, 2)
    elif tag == "h3":
        add_heading(doc, node, 3)
    elif tag == "p":
        class_names = set(node.get("class", []))
        if "caption" in class_names or "appendix-note" in class_names:
            add_caption(doc, text_content(node), figure=False)
        else:
            add_paragraph_from_node(doc, node)
    elif tag == "aside":
        add_callout(doc, text_content(node))
    elif tag == "table":
        add_table(doc, node)
    elif tag == "figure":
        add_figure(doc, project, node)
    elif tag == "ol":
        add_list(doc, node, ordered=True, numbering_ids=numbering_ids)
    elif tag == "ul":
        add_list(doc, node, ordered=False, numbering_ids=numbering_ids)
    elif tag == "img":
        src = str(node.get("src", ""))
        image_path = resolve_image(project, src)
        if image_path:
            doc.add_picture(str(image_path), width=Inches(6.1))
    else:
        for child in iter_meaningful_children(node):
            render_node(doc, project, child, numbering_ids)


def chapter_paths(project: Path) -> list[Path]:
    manifest_path = project / "reports" / "report_assembly_manifest.json"
    chapters_dir = project / "reports" / "chapters"
    if manifest_path.exists():
        manifest = json.loads(read_text(manifest_path))
        ordered = manifest.get("assembled_chapters") or manifest.get("chapters")
        if isinstance(ordered, list) and ordered:
            paths = [project / str(item) if str(item).startswith("reports/") else chapters_dir / str(item) for item in ordered]
        else:
            paths = sorted(chapters_dir.glob("ch*.html"))
    else:
        paths = sorted(chapters_dir.glob("ch*.html"))
    missing = [path.as_posix() for path in paths if not path.exists()]
    if missing:
        raise FileNotFoundError("chapter file(s) missing: " + " | ".join(missing))
    if not paths:
        raise FileNotFoundError("no chapter files found under reports/chapters")
    return paths


def add_chapters(doc: Document, project: Path, paths: list[Path], numbering_ids: dict[str, int]) -> None:
    for path in paths:
        soup = BeautifulSoup(read_text(path), "html.parser")
        roots = [child for child in iter_meaningful_children(soup)]
        for root in roots:
            render_node(doc, project, root, numbering_ids)


def csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return [{str(k or "").strip(): str(v or "").strip() for k, v in row.items()} for row in csv.DictReader(handle)]


def add_reference_appendices(doc: Document, project: Path) -> None:
    source_rows = [
        row for row in csv_rows(project / "references" / "source_link_register.csv")
        if row.get("title") or row.get("publisher") or row.get("official_url") or row.get("url")
    ]
    if source_rows:
        p = doc.add_paragraph(style="Heading 1")
        p.add_run("참고자료")
        table = doc.add_table(rows=1, cols=4)
        set_table_width(table)
        set_table_borders(table)
        headers = ["No.", "자료명", "발행기관", "원문"]
        widths = [780, 4200, 2100, 2280]
        for index, label in enumerate(headers):
            cell = table.rows[0].cells[index]
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            set_cell_shading(cell, COLOR_SOFT)
            add_text(cell.paragraphs[0], label, size=9, color=COLOR_DARK, bold=True)
        for index, row in enumerate(source_rows, start=1):
            cells = table.add_row().cells
            values = [
                str(index),
                row.get("title", ""),
                row.get("publisher", ""),
                "원문 링크" if row.get("official_url") or row.get("url") else "",
            ]
            for col, value in enumerate(values):
                set_cell_width(cells[col], widths[col])
                set_cell_margins(cells[col])
                add_text(cells[col].paragraphs[0], value, size=8.8, color=COLOR_INK)
        add_caption(doc, "자료: 출처 링크 등록표와 source records. 근거 데이터: 참고자료 목록.")

    data_files = sorted(
        path for path in (project / "data_sources").glob("*")
        if path.is_file() and path.suffix.lower() in {".csv", ".xlsx", ".xls", ".tsv"} and path.name != "visual_plan.csv"
    )
    if data_files:
        p = doc.add_paragraph(style="Heading 1")
        p.add_run("부록. 분석 데이터 목록")
        table = doc.add_table(rows=1, cols=4)
        set_table_width(table)
        set_table_borders(table)
        headers = ["No.", "데이터셋", "파일명", "크기(bytes)"]
        widths = [780, 3600, 3600, 1380]
        for index, label in enumerate(headers):
            cell = table.rows[0].cells[index]
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            set_cell_shading(cell, COLOR_SOFT)
            add_text(cell.paragraphs[0], label, size=9, color=COLOR_DARK, bold=True)
        for index, path in enumerate(data_files, start=1):
            cells = table.add_row().cells
            values = [str(index), path.stem.replace("_", " "), path.name, str(path.stat().st_size)]
            for col, value in enumerate(values):
                set_cell_width(cells[col], widths[col])
                set_cell_margins(cells[col])
                add_text(cells[col].paragraphs[0], value, size=8.8, color=COLOR_INK)
        add_caption(doc, "자료: visual plan과 data_sources 등록 파일. 근거 데이터: 분석 데이터 목록.")


def docx_structure(path: Path) -> dict[str, object]:
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
    return {
        "valid_package": True,
        "parts": len(names),
        "has_document": "word/document.xml" in names,
        "has_styles": "word/styles.xml" in names,
        "has_numbering": "word/numbering.xml" in names,
        "has_headers": any(name.startswith("word/header") for name in names),
        "has_footers": any(name.startswith("word/footer") for name in names),
        "has_media": any(name.startswith("word/media/") for name in names),
        "has_footnotes_or_endnotes": "word/footnotes.xml" in names or "word/endnotes.xml" in names,
    }


def find_soffice() -> str | None:
    program_files = [os.environ.get("ProgramFiles", ""), os.environ.get("ProgramFiles(x86)", "")]
    libreoffice_candidates = [
        str(Path(root) / "LibreOffice" / "program" / "soffice.exe")
        for root in program_files
        if root
    ]
    candidates = [
        shutil.which("soffice"),
        *libreoffice_candidates,
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def find_pdftoppm() -> str | None:
    bundled = (
        Path.home()
        / ".cache"
        / "codex-runtimes"
        / "codex-primary-runtime"
        / "dependencies"
        / "native"
        / "poppler"
        / "Library"
        / "bin"
        / "pdftoppm.exe"
    )
    candidates = [
        str(bundled),
        shutil.which("pdftoppm"),
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def render_preview(docx_path: Path, output_dir: Path) -> dict[str, object]:
    docx_path = docx_path.resolve()
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    soffice = find_soffice()
    if not soffice:
        return {"status": "not_run", "reason": "LibreOffice soffice was not found"}
    run_key = datetime.now().strftime("%H%M%S%f")
    render_dir = (Path(".local_state") / "docx_render" / f"r_{run_key}").resolve()
    render_dir.mkdir(parents=True, exist_ok=True)
    profile = (render_dir / "lo_profile").resolve()
    profile.mkdir(parents=True, exist_ok=True)
    render_input = render_dir / "input.docx"
    shutil.copy2(docx_path, render_input)
    command = [
        soffice,
        "--headless",
        "-env:UserInstallation=" + "file:" + "/" * 3 + profile.as_posix(),
        "--convert-to",
        "pdf",
        "--outdir",
        str(render_dir),
        str(render_input),
    ]
    try:
        proc = subprocess.run(
            command,
            text=True,
            capture_output=True,
            timeout=60,
        )
    except subprocess.TimeoutExpired as exc:
        return {
            "status": "failed",
            "reason": "LibreOffice conversion timed out",
            "command": command,
            "timeout_seconds": exc.timeout,
        }
    pdf = render_dir / f"{render_input.stem}.pdf"
    result: dict[str, object] = {
        "status": "pdf_created" if pdf.exists() else "failed",
        "returncode": proc.returncode,
        "stdout": proc.stdout[-1000:],
        "stderr": proc.stderr[-1000:],
        "pdf": "",
        "pages": [],
    }
    pdftoppm = find_pdftoppm()
    if pdf.exists() and pdftoppm:
        prefix = render_dir / "page"
        subprocess.run([pdftoppm, "-png", "-r", "150", str(pdf), str(prefix)], check=False)
        pages = sorted(render_dir.glob("page-*.png"))
        copied_pages: list[str] = []
        copied_pdf = output_dir / f"{docx_path.stem}.pdf"
        shutil.copy2(pdf, copied_pdf)
        for page in pages:
            target = output_dir / f"{docx_path.stem}_{page.name}"
            shutil.copy2(page, target)
            copied_pages.append(target.as_posix())
        result["status"] = "rendered_png" if pages else result["status"]
        result["pdf"] = copied_pdf.as_posix()
        result["pages"] = copied_pages
    elif pdf.exists():
        copied_pdf = output_dir / f"{docx_path.stem}.pdf"
        shutil.copy2(pdf, copied_pdf)
        result["pdf"] = copied_pdf.as_posix()
        result["png_note"] = "pdftoppm was not found; only PDF render was created"
    return result


def export_docx(project: Path, cover_data_path: Path, output_path: Path, render: bool) -> dict[str, object]:
    cover_data = json.loads(read_text(cover_data_path))
    chapters = chapter_paths(project)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_document(doc)
    list_style_presets = load_list_style_presets()
    numbering_ids = configure_list_numbering(doc, list_style_presets)
    set_running_header_footer(doc, cover_data)
    add_cover(doc, cover_data)
    add_chapters(doc, project, chapters, numbering_ids)
    add_reference_appendices(doc, project)
    doc.save(output_path)

    checks_dir = project / "reports" / "export_checks"
    rel_output = output_path.relative_to(project).as_posix()
    structure = docx_structure(output_path)
    payload: dict[str, object] = {
        "project": project.name,
        "generated_at_kst": now_kst(),
        "status": "structure_checked",
        "export_type": "native_docx",
        "preset": "standard_business_brief with report_orchestra_a4_native_docx override",
        "source_cover_data": cover_data_path.relative_to(project).as_posix(),
        "source_chapters": [path.relative_to(project).as_posix() for path in chapters],
        "output": rel_output,
        "sha256": sha256_file(output_path),
        "structure": structure,
        "list_style_presets": sorted(numbering_ids),
        "limitations": [
            "HTML preview fidelity is not used as the export mechanism.",
            "Native multi-level list presets are generated for DOCX, but exact visual parity with Word/Google Docs HTML import still requires export verification.",
            "Native Word field codes, SEQ captions, editable charts, and mixed-orientation sections are not generated by this tool.",
            "HTML footnote-like sections are preserved as reader-facing sections unless a separate native footnote workflow is added.",
        ],
    }
    if render:
        render_dir = checks_dir / output_path.stem
        render_result = render_preview(output_path, render_dir)
        payload["render_check"] = render_result
        if render_result.get("status") == "rendered_png":
            payload["status"] = "render_verified_needs_visual_review"
    write_json(checks_dir / "docx_structure_check.json", payload)
    return payload


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser(description="Export report sources to a native Word DOCX artifact.")
    parser.add_argument("--project", required=True, help="Project folder name under 00_사용자_작업공간")
    parser.add_argument("--cover-data", default="reports/cover.data.json", help="Cover data JSON path relative to project")
    parser.add_argument("--output", default=DEFAULT_OUTPUT.as_posix(), help="DOCX output path relative to project")
    parser.add_argument("--render-preview", action="store_true", help="Render the DOCX to PDF/PNGs under reports/export_checks/")
    args = parser.parse_args()

    project = PROJECT_ROOT / args.project
    if not project.exists():
        print(json.dumps({"error": f"project not found: {args.project}"}, ensure_ascii=False, indent=2))
        return 2
    cover_data_path = project / args.cover_data
    if not cover_data_path.exists():
        print(json.dumps({"error": f"cover data not found: {args.cover_data}"}, ensure_ascii=False, indent=2))
        return 2
    output_path = project / args.output
    try:
        payload = export_docx(project, cover_data_path, output_path, args.render_preview)
    except Exception as exc:
        print(json.dumps({"error": str(exc)}, ensure_ascii=False, indent=2))
        return 1
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
